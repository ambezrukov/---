#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Анализатор документов с поддержкой различных форматов
Поддерживает PDF, DOCX, RTF, TXT, изображения и архивы

Версия: 2.0.0 "ZIP Master"
Дата выпуска: 20 июля 2025
Статус: Стабильная версия
"""

import os
import sys
import json
import hashlib
import threading
import queue
import time
from datetime import datetime, timedelta
from pathlib import Path
import traceback
import zipfile
import shutil
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Dict, List, Tuple, Optional

# GUI imports
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkinter.scrolledtext import ScrolledText

# Document processing imports
try:
    import pytesseract
    from PIL import Image
    import pdfplumber
    from docx import Document
    
    # Настройка пути к Tesseract (та же логика, что и в основной программе)
    tesseract_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        "/usr/bin/tesseract",
        "/usr/local/bin/tesseract"
    ]
    
    tesseract_found = False
    for path in tesseract_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            print(f"   ✅ Tesseract найден: {path}")
            tesseract_found = True
            break
    
    if not tesseract_found:
        # Попробуем найти через команду
        try:
            import subprocess
            result = subprocess.run(['tesseract', '--version'], 
                                  capture_output=True, text=True, timeout=5)
            if result.returncode == 0:
                print("   ✅ Tesseract найден через PATH")
                tesseract_found = True
        except:
            pass
    
    if not tesseract_found:
        # Попробуем найти через команду
        try:
            import subprocess
            result = subprocess.run(['tesseract', '--version'], 
                                  capture_output=True, text=True, timeout=5)
            if result.returncode == 0:
                print("✅ Tesseract найден через PATH")
                tesseract_found = True
        except:
            pass
    
    if not tesseract_found:
        print("❌ Tesseract не найден. Установите Tesseract OCR для обработки отсканированных документов.")

    TESSERACT_AVAILABLE = tesseract_found

except ImportError as e:
    TESSERACT_AVAILABLE = False
    print(f"Предупреждение: {e}")

# Try to import textract for .doc and .rtf files
try:
    import textract
    TEXTTRACT_AVAILABLE = True
except ImportError:
    TEXTTRACT_AVAILABLE = False
    print("Предупреждение: textract не установлен. Файлы .doc/.rtf не будут поддерживаться.")

class ErrorLogger:
    """Система логирования ошибок"""
    
    def __init__(self, log_dir: str = "logs"):
        self.log_dir = Path(log_dir)
        self.log_dir.mkdir(exist_ok=True)
        self.error_log_file = self.log_dir / f"errors_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
        self.session_log_file = self.log_dir / f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
        
        # Записываем заголовок в файлы логов
        self._write_log_header(self.error_log_file, "ЛОГ ОШИБОК")
        self._write_log_header(self.session_log_file, "ЛОГ СЕССИИ")
    
    def _write_log_header(self, log_file: Path, title: str):
        """Запись заголовка в лог файл"""
        try:
            with open(log_file, 'w', encoding='utf-8') as f:
                f.write(f"{'='*60}\n")
                f.write(f"{title}\n")
                f.write(f"{'='*60}\n")
                f.write(f"Дата и время: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Система: {os.name}\n")
                f.write(f"Python: {sys.version}\n")
                f.write(f"{'='*60}\n\n")
        except Exception as e:
            print(f"Ошибка создания лог файла: {e}")
    
    def log_error(self, filepath: str, error: str, error_type: str = "ОБРАБОТКА"):
        """Запись ошибки в лог"""
        try:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            with open(self.error_log_file, 'a', encoding='utf-8') as f:
                f.write(f"[{timestamp}] {error_type}: {filepath}\n")
                f.write(f"Ошибка: {error}\n")
                f.write(f"{'-'*40}\n")
        except Exception as e:
            print(f"Ошибка записи в лог: {e}")
    
    def log_session(self, message: str, level: str = "INFO"):
        """Запись сообщения сессии в лог"""
        try:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            with open(self.session_log_file, 'a', encoding='utf-8') as f:
                f.write(f"[{timestamp}] {level}: {message}\n")
        except Exception as e:
            print(f"Ошибка записи в лог сессии: {e}")
    
    def get_error_summary(self) -> Dict[str, int]:
        """Получение сводки ошибок"""
        error_counts = {}
        try:
            if self.error_log_file.exists():
                with open(self.error_log_file, 'r', encoding='utf-8') as f:
                    for line in f:
                        if 'ОШИБКА:' in line:
                            error_type = line.split('ОШИБКА:')[1].strip()
                            error_counts[error_type] = error_counts.get(error_type, 0) + 1
        except Exception as e:
            print(f"Ошибка чтения сводки: {e}")
        return error_counts
    
    def export_errors_to_txt(self, output_file: str = None) -> str:
        """Экспорт ошибок в текстовый файл"""
        if output_file is None:
            output_file = f"errors_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        
        try:
            if self.error_log_file.exists():
                shutil.copy2(self.error_log_file, output_file)
                return output_file
        except Exception as e:
            print(f"Ошибка экспорта: {e}")
        return ""

class DependencyChecker:
    """Проверка и установка зависимостей"""
    
    @staticmethod
    def check_tesseract():
        """Проверка установки Tesseract OCR"""
        try:
            # Попытка найти Tesseract в стандартных местах
            possible_paths = [
                r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                '/usr/bin/tesseract',
                '/usr/local/bin/tesseract'
            ]
            
            for path in possible_paths:
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
                    print(f"✅ Tesseract найден: {path}")
                    return True, path
            
            # Попытка найти через команду
            import subprocess
            result = subprocess.run(['tesseract', '--version'], 
                                  capture_output=True, text=True)
            if result.returncode == 0:
                return True, "tesseract"
                
            return False, "Tesseract не найден"
        except Exception as e:
            return False, str(e)
    
    @staticmethod
    def get_missing_dependencies():
        """Получить список отсутствующих зависимостей"""
        missing = []
        
        if not TESSERACT_AVAILABLE:
            missing.append("pytesseract")
        
        if not TEXTTRACT_AVAILABLE:
            missing.append("textract")
        
        tesseract_ok, tesseract_msg = DependencyChecker.check_tesseract()
        if not tesseract_ok:
            missing.append(f"Tesseract OCR: {tesseract_msg}")
        
        return missing

class CacheManager:
    """Управление кэшем результатов"""
    
    def __init__(self, cache_dir: str = ".cache"):
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(exist_ok=True)
        self.cache_file = self.cache_dir / "file_cache.json"
        self.load_cache()
    
    def load_cache(self):
        """Загрузка кэша из файла"""
        try:
            if self.cache_file.exists():
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    self.cache = json.load(f)
            else:
                self.cache = {}
        except Exception:
            self.cache = {}
    
    def save_cache(self):
        """Сохранение кэша в файл"""
        try:
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.cache, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Ошибка сохранения кэша: {e}")
    
    def get_file_hash(self, filepath: str) -> str:
        """Получить хеш файла для проверки изменений"""
        try:
            stat = os.stat(filepath)
            return f"{stat.st_mtime}_{stat.st_size}"
        except Exception:
            return ""
    
    def get_cached_text(self, filepath: str) -> Optional[str]:
        """Получить текст из кэша"""
        file_hash = self.get_file_hash(filepath)
        cached_data = self.cache.get(filepath)
        
        if cached_data and cached_data.get('hash') == file_hash:
            return cached_data.get('text', "")
        return None
    
    def cache_text(self, filepath: str, text: str):
        """Сохранить текст в кэш"""
        file_hash = self.get_file_hash(filepath)
        self.cache[filepath] = {
            'hash': file_hash,
            'text': text,
            'timestamp': datetime.now().isoformat()
        }
    
    def clear_cache(self):
        """Очистить кэш"""
        self.cache.clear()
        if self.cache_file.exists():
            self.cache_file.unlink()

class DocumentProcessor:
    """Обработчик документов с поддержкой различных форматов"""
    
    SUPPORTED_EXTENSIONS = {
        'images': ('.jpg', '.jpeg', '.png', '.bmp', '.tiff'),
        'documents': ('.docx', '.doc', '.rtf'),
        'pdfs': ('.pdf',),
        'text': ('.txt', '.md'),
        'archives': ('.zip', '.rar', '.7z')
    }
    
    def __init__(self, cache_manager: CacheManager):
        self.cache_manager = cache_manager
        self._cancelled = False
        self.error_log = []
    
    def extract_text(self, filepath: str) -> Tuple[str, bool]:
        """Извлечение текста из файла"""
        # Проверяем кэш
        cached_text = self.cache_manager.get_cached_text(filepath)
        if cached_text is not None:
            return cached_text, True
        
        ext = Path(filepath).suffix.lower()
        text = ""
        success = False
        
        try:
            if ext in self.SUPPORTED_EXTENSIONS['pdfs']:
                text = self._extract_from_pdf(filepath)
            elif ext in self.SUPPORTED_EXTENSIONS['images']:
                text = self._extract_from_image(filepath)
            elif ext == '.docx':
                text = self._extract_from_docx(filepath)
            elif ext in ('.doc', '.rtf'):
                text = self._extract_from_legacy_doc(filepath)
            elif ext in self.SUPPORTED_EXTENSIONS['text']:
                text = self._extract_from_text(filepath)
            elif ext in self.SUPPORTED_EXTENSIONS['archives']:
                text = self._extract_from_archive(filepath)
            
            if text and text.strip():
                success = True
                # Сохраняем в кэш
                self.cache_manager.cache_text(filepath, text)
            else:
                # Если текст пустой, записываем в лог
                error_msg = f"Файл пустой или не содержит текста: {filepath}"
                self.error_log.append(error_msg)
                
        except Exception as e:
            error_msg = f"Ошибка при обработке {filepath}: {str(e)}"
            self.error_log.append(error_msg)
            
            # Записываем в лог файл для диагностики
            import traceback
            full_error = f"Полная ошибка для {filepath}:\n{traceback.format_exc()}"
            print(full_error)  # Выводим в консоль для отладки
        
        return text, success
    
    def set_cancelled(self, cancelled: bool):
        """Установка флага отмены обработки"""
        self._cancelled = cancelled
    
    def _extract_from_pdf(self, filepath: str) -> str:
        """Извлечение текста из PDF"""
        try:
            print(f"Попытка обработки PDF: {filepath}")
            
            # Проверяем размер файла
            file_size = os.path.getsize(filepath)
            print(f"Размер файла: {file_size} байт")
            if file_size == 0:
                raise Exception("Файл пустой (0 байт)")
            
            # Метод 1: pdfplumber
            texts = []
            try:
                with pdfplumber.open(filepath) as pdf:
                    print(f"PDF открыт через pdfplumber, страниц: {len(pdf.pages)}")
                    
                    # Обрабатываем все страницы, но с проверкой прерывания
                    pages_to_process = len(pdf.pages)
                    print(f"📄 Обрабатываем все {pages_to_process} страниц...")
                    
                    # Проверяем, не слишком ли большой файл
                    if pages_to_process > 200:
                        print(f"⚠️ ВНИМАНИЕ: Файл содержит {pages_to_process} страниц. Обработка может занять значительное время.")
                        texts.append(f"⚠️ ВНИМАНИЕ: Файл содержит {pages_to_process} страниц. Обработка может занять значительное время.\n")
                    
                    # Проверяем, есть ли защита
                    if hasattr(pdf, 'metadata') and pdf.metadata:
                        print(f"Метаданные PDF: {pdf.metadata}")
                    
                    for i, page in enumerate(pdf.pages):
                        try:
                            # Проверяем, не было ли прерывания
                            if hasattr(self, '_cancelled') and self._cancelled:
                                print("Обработка прервана пользователем")
                                break
                                
                            # Показываем прогресс каждые 10 страниц или для больших файлов
                            if i % 10 == 0 or i == len(pdf.pages) - 1:
                                print(f"📄 Обработка страницы {i+1}/{len(pdf.pages)}...")
                                
                            text = page.extract_text()
                            if text and text.strip():
                                texts.append(f"--- СТРАНИЦА {i+1} ---\n{text}")
                                print(f"Страница {i+1}: извлечено {len(text)} символов")
                            else:
                                print(f"Страница {i+1}: текст пустой")
                                # Проверяем, есть ли изображения на странице
                                if page.images:
                                    print(f"Страница {i+1}: найдено {len(page.images)} изображений")
                        except Exception as e:
                            print(f"Ошибка страницы {i+1}: {str(e)}")
            except Exception as e:
                print(f"pdfplumber ошибка: {str(e)}")
                # Проверяем, не защищен ли файл паролем
                if "password" in str(e).lower() or "encrypted" in str(e).lower():
                    raise Exception(f"PDF файл защищен паролем: {str(e)}")
            
            # Метод 2: PyMuPDF (fitz) - более мощная библиотека для отсканированных документов
            if not texts:
                print("pdfplumber не дал результатов, пробуем PyMuPDF...")
                try:
                    import fitz  # PyMuPDF
                    doc = fitz.open(filepath)
                    print(f"PyMuPDF: страниц: {len(doc)}")
                    
                    # Обрабатываем все страницы, но с проверкой прерывания
                    pages_to_process = len(doc)
                    print(f"📄 Обрабатываем все {pages_to_process} страниц через PyMuPDF...")
                    
                    # Проверяем, не слишком ли большой файл
                    if pages_to_process > 200:
                        print(f"⚠️ ВНИМАНИЕ: Файл содержит {pages_to_process} страниц. Обработка может занять значительное время.")
                        texts.append(f"⚠️ ВНИМАНИЕ: Файл содержит {pages_to_process} страниц. Обработка может занять значительное время.\n")
                    
                    for i in range(pages_to_process):
                        try:
                            # Проверяем, не было ли прерывания
                            if hasattr(self, '_cancelled') and self._cancelled:
                                print("Обработка прервана пользователем")
                                break
                                
                            page = doc.load_page(i)
                            text = page.get_text()
                            if text and text.strip():
                                texts.append(f"--- СТРАНИЦА {i+1} (PyMuPDF) ---\n{text}")
                                print(f"PyMuPDF страница {i+1}: извлечено {len(text)} символов")
                            else:
                                print(f"PyMuPDF страница {i+1}: текст пустой")
                                # Проверяем, есть ли изображения
                                image_list = page.get_images()
                                if image_list:
                                    print(f"PyMuPDF страница {i+1}: найдено {len(image_list)} изображений")
                        except Exception as e:
                            print(f"PyMuPDF ошибка страницы {i+1}: {str(e)}")
                    
                    doc.close()
                except ImportError:
                    print("PyMuPDF не установлен")
                except Exception as e:
                    print(f"PyMuPDF ошибка: {str(e)}")
            
            # Метод 3: PyPDF2 (если предыдущие не дали результатов)
            if not texts:
                print("Пробуем PyPDF2...")
                try:
                    import PyPDF2
                    with open(filepath, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        print(f"PyPDF2: страниц: {len(pdf_reader.pages)}")
                        
                        # Обрабатываем все страницы, но с проверкой прерывания
                        pages_to_process = len(pdf_reader.pages)
                        print(f"📄 Обрабатываем все {pages_to_process} страниц через PyPDF2...")
                        
                        # Проверяем, не слишком ли большой файл
                        if pages_to_process > 200:
                            print(f"⚠️ ВНИМАНИЕ: Файл содержит {pages_to_process} страниц. Обработка может занять значительное время.")
                            texts.append(f"⚠️ ВНИМАНИЕ: Файл содержит {pages_to_process} страниц. Обработка может занять значительное время.\n")
                        
                        # Проверяем защиту
                        if pdf_reader.is_encrypted:
                            raise Exception("PDF файл зашифрован")
                        
                        for i, page in enumerate(pdf_reader.pages):
                            try:
                                # Проверяем, не было ли прерывания
                                if hasattr(self, '_cancelled') and self._cancelled:
                                    print("Обработка прервана пользователем")
                                    break
                                    
                                text = page.extract_text()
                                if text and text.strip():
                                    texts.append(f"--- СТРАНИЦА {i+1} (PyPDF2) ---\n{text}")
                                    print(f"PyPDF2 страница {i+1}: извлечено {len(text)} символов")
                                else:
                                    print(f"PyPDF2 страница {i+1}: текст пустой")
                            except Exception as e:
                                print(f"PyPDF2 ошибка страницы {i+1}: {str(e)}")
                except ImportError:
                    print("PyPDF2 не установлен")
                except Exception as e:
                    print(f"PyPDF2 ошибка: {str(e)}")
                    if "encrypted" in str(e).lower():
                        raise Exception(f"PDF файл зашифрован: {str(e)}")
            
            # Метод 4: OCR (если все еще нет текста и Tesseract доступен)
            if not texts:
                print("Текстовый слой пустой, пробуем OCR...")
                
                # Проверяем доступность Tesseract
                if not TESSERACT_AVAILABLE:
                    print("Tesseract OCR не установлен. Для обработки отсканированных документов установите Tesseract.")
                    raise Exception("Tesseract OCR не установлен. Установите Tesseract для обработки отсканированных документов.")
                
                try:
                    # Используем pdf2image для конвертации
                    from pdf2image import convert_from_path
                    print("Конвертируем PDF в изображения...")
                    
                    # Проверяем, установлен ли poppler
                    poppler_paths = [
                        r"C:\poppler\bin",
                        r"C:\Program Files\poppler\bin", 
                        r"C:\Program Files (x86)\poppler\bin",
                        os.path.join(os.path.expanduser("~"), "poppler", "bin"),
                        r"C:\Users\Home\poppler-24.08.0\Library\bin"
                    ]
                    
                    poppler_found = False
                    poppler_path = None
                    for path in poppler_paths:
                        if os.path.exists(os.path.join(path, "pdftoppm.exe")):
                            print(f"Poppler найден: {path}")
                            poppler_found = True
                            poppler_path = path
                            break
                    
                    if not poppler_found:
                        # Проверяем через команду
                        try:
                            import subprocess
                            result = subprocess.run(['pdftoppm', '-h'], 
                                                  capture_output=True, text=True, timeout=5)
                            if result.returncode == 0:
                                print("Poppler найден через PATH")
                                poppler_found = True
                        except:
                            pass
                    
                    if not poppler_found:
                        raise Exception("Poppler не найден")
                    
                    # Обрабатываем все страницы через OCR, но с возможностью прерывания
                    print("🔄 Конвертируем все страницы PDF в изображения для OCR...")
                    
                    # Используем найденный путь к poppler
                    if poppler_path:
                        images = convert_from_path(filepath, dpi=300, poppler_path=poppler_path)
                    else:
                        images = convert_from_path(filepath, dpi=300)
                        
                    print(f"📷 Конвертировано {len(images)} страниц в изображения")
                    
                    if len(images) > 200:
                        print(f"⚠️ ВНИМАНИЕ: Файл содержит {len(images)} страниц. OCR может занять значительное время.")
                        texts.append(f"⚠️ ВНИМАНИЕ: Файл содержит {len(images)} страниц. OCR может занять значительное время.\n")
                    
                    for i, image in enumerate(images):
                        try:
                            # Проверяем, не было ли прерывания
                            if hasattr(self, '_cancelled') and self._cancelled:
                                print("OCR прерван пользователем")
                                break
                                
                            print(f"OCR страницы {i+1}/{len(images)}...")
                            
                            # OCR с русским и английским языками
                            ocr_text = pytesseract.image_to_string(image, lang='rus+eng', config='--psm 6')
                            if ocr_text and ocr_text.strip():
                                texts.append(f"--- СТРАНИЦА {i+1} (OCR) ---\n{ocr_text}")
                                print(f"OCR страница {i+1}: извлечено {len(ocr_text)} символов")
                            else:
                                print(f"OCR страница {i+1}: текст не найден")
                        except Exception as ocr_error:
                            print(f"OCR ошибка страницы {i+1}: {str(ocr_error)}")
                            
                except Exception as convert_error:
                    error_msg = str(convert_error).lower()
                    if "poppler" in error_msg or "unable to get page count" in error_msg:
                        print("❌ Poppler не установлен или не найден в PATH")
                        print("💡 Для OCR требуется установить poppler-utils")
                        print("   Скачайте с https://github.com/oschwartz10612/poppler-windows/releases")
                        print("   Распакуйте в папку пользователя и добавьте в PATH")
                        raise Exception("Для OCR требуется установить poppler-utils. Скачайте с https://github.com/oschwartz10612/poppler-windows/releases")
                    else:
                        raise Exception(f"Ошибка конвертации PDF: {str(convert_error)}")
                            
                except ImportError:
                    print("pdf2image не установлен")
                    raise Exception("pdf2image не установлен. Установите: pip install pdf2image")
                except Exception as e:
                    print(f"OCR ошибка: {str(e)}")
                    raise e
            
            result = "\n\n".join(texts) if texts else ""
            print(f"Итоговый результат: {len(result)} символов")
            
            if not result.strip():
                raise Exception("Не удалось извлечь текст из PDF. Возможные причины: файл содержит только изображения без текстового слоя, или требуется установка Tesseract OCR для распознавания отсканированных документов.")
                
            return result
                
        except Exception as e:
            print(f"Критическая ошибка PDF {filepath}: {str(e)}")
            raise Exception(f"Ошибка открытия PDF {filepath}: {str(e)}")
    
    def _extract_from_image(self, filepath: str) -> str:
        """Извлечение текста из изображения (OCR)"""
        if not TESSERACT_AVAILABLE:
            raise ImportError("pytesseract не установлен")
        
        image = Image.open(filepath)
        return pytesseract.image_to_string(image, lang='rus+eng')
    
    def _extract_from_docx(self, filepath: str) -> str:
        """Извлечение текста из DOCX"""
        doc = Document(filepath)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        return "\n".join(texts)
    
    def _extract_from_legacy_doc(self, filepath: str) -> str:
        """Извлечение текста из DOC/RTF"""
        if not TEXTTRACT_AVAILABLE:
            raise ImportError("Файлы .doc/.rtf не поддерживаются. Для их обработки требуется установка библиотеки textract, которая не совместима с Python 3.13.")
        return textract.process(filepath).decode('utf-8')
    
    def _extract_from_text(self, filepath: str) -> str:
        """Извлечение текста из текстовых файлов"""
        encodings = ['utf-8', 'cp1251', 'latin-1']
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError(f"Не удалось декодировать {filepath}")

    def _extract_from_pdf_simple(self, filepath: str) -> str:
        """Упрощенная обработка PDF через OCR для архивов"""
        try:
            print(f"Упрощенная обработка PDF: {filepath}")
            
            # Проверяем размер файла
            file_size = os.path.getsize(filepath)
            print(f"Размер файла: {file_size} байт")
            if file_size == 0:
                raise Exception("Файл пустой (0 байт)")
            

            
            # Сначала пробуем извлечь текст без OCR
            texts = []
            try:
                with pdfplumber.open(filepath) as pdf:
                    print(f"PDF открыт через pdfplumber, страниц: {len(pdf.pages)}")
                    
                    for i, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        if text and text.strip():
                            texts.append(f"--- СТРАНИЦА {i+1} ---\n{text}")
                            print(f"Страница {i+1}: извлечено {len(text)} символов")
                        else:
                            print(f"Страница {i+1}: текст пустой")
            except Exception as e:
                print(f"pdfplumber ошибка: {str(e)}")
            
            # Если текста нет, используем OCR
            if not texts:
                print("Текстовый слой пустой, пробуем OCR...")
                
                # Проверяем доступность Tesseract
                if not TESSERACT_AVAILABLE:
                    print("Tesseract OCR не установлен.")
                    raise Exception("Tesseract OCR не установлен.")
                
                try:
                    # Используем pdf2image для конвертации
                    from pdf2image import convert_from_path
                    print("Конвертируем PDF в изображения...")
                    
                    # Ищем poppler
                    poppler_paths = [
                        r"C:\poppler\bin",
                        r"C:\Program Files\poppler\bin", 
                        r"C:\Program Files (x86)\poppler\bin",
                        os.path.join(os.path.expanduser("~"), "poppler", "bin"),
                        r"C:\Users\Home\poppler-24.08.0\Library\bin"
                    ]
                    
                    poppler_found = False
                    poppler_path = None
                    for path in poppler_paths:
                        if os.path.exists(os.path.join(path, "pdftoppm.exe")):
                            print(f"Poppler найден: {path}")
                            poppler_found = True
                            poppler_path = path
                            break
                    
                    if not poppler_found:
                        raise Exception("Poppler не найден")
                    
                    # Конвертируем
                    if poppler_path:
                        images = convert_from_path(filepath, dpi=300, poppler_path=poppler_path)
                    else:
                        images = convert_from_path(filepath, dpi=300)
                    
                    print(f"Получено {len(images)} изображений")
                    
                    # OCR каждое изображение
                    import subprocess
                    tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
                    
                    for i, image in enumerate(images):
                        print(f"OCR страница {i+1}...")
                        
                        # Сохраняем изображение во временный файл
                        import tempfile
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as img_file:
                            image.save(img_file.name, 'PNG')
                            img_path = img_file.name
                        
                        try:
                            # OCR через Tesseract
                            result = subprocess.run([tesseract_path, img_path, 'stdout', '-l', 'rus+eng'], 
                                                  capture_output=True, timeout=30)
                            
                            if result.returncode == 0:
                                try:
                                    text = result.stdout.decode('utf-8').strip()
                                except UnicodeDecodeError:
                                    try:
                                        text = result.stdout.decode('cp1251').strip()
                                    except:
                                        text = result.stdout.decode('utf-8', errors='ignore').strip()
                                
                                if text:
                                    texts.append(f"--- СТРАНИЦА {i+1} (OCR) ---\n{text}")
                                    print(f"Страница {i+1}: OCR извлечено {len(text)} символов")
                                else:
                                    print(f"Страница {i+1}: OCR не дал результатов")
                            else:
                                print(f"Ошибка OCR страницы {i+1}: {result.stderr}")
                        
                        finally:
                            # Удаляем временное изображение
                            try:
                                os.unlink(img_path)
                            except:
                                pass
                
                except ImportError:
                    print("pdf2image не установлен")
                    raise Exception("pdf2image не установлен")
                except Exception as e:
                    print(f"Ошибка OCR: {e}")
                    raise Exception(f"Ошибка OCR: {e}")
            
            result = "\n\n".join(texts) if texts else ""
            print(f"Итоговый результат: {len(result)} символов")
            return result
            
        except Exception as e:
            raise Exception(f"Ошибка обработки PDF: {str(e)}")
    
    def _extract_from_archive(self, filepath: str) -> str:
        """Извлечение текста из архивов (ZIP/RAR)"""
        ext = Path(filepath).suffix.lower()
        all_texts = []
        
        try:
            if ext == '.zip':
                print(f"Попытка обработки ZIP: {filepath}")
                
                # Проверяем размер файла
                file_size = os.path.getsize(filepath)
                print(f"Размер ZIP файла: {file_size} байт")
                if file_size == 0:
                    raise Exception("ZIP файл пустой (0 байт)")
                
                with zipfile.ZipFile(filepath, 'r') as zip_file:
                    # Проверяем, не поврежден ли архив
                    try:
                        zip_file.testzip()
                        print("ZIP файл не поврежден")
                    except Exception as test_error:
                        print(f"ZIP файл поврежден: {str(test_error)}")
                        raise Exception(f"ZIP файл поврежден: {str(test_error)}")
                    
                    # Получаем список файлов в архиве
                    file_list = zip_file.namelist()
                    print(f"ZIP содержит {len(file_list)} файлов: {file_list}")
                    
                    if not file_list:
                        raise Exception("ZIP архив пустой")
                    
                    # Обрабатываем только текстовые файлы и документы
                    supported_extensions = ('.txt', '.md', '.pdf', '.docx', '.doc', '.rtf')
                    found_supported_files = False
                    processed_files = 0
                    
                    for file_name in file_list:
                        if any(file_name.lower().endswith(ext_suffix) for ext_suffix in supported_extensions):
                            found_supported_files = True
                            try:
                                print(f"Обрабатываем файл в архиве: {file_name}")
                                
                                # Читаем файл из архива
                                with zip_file.open(file_name) as file_in_archive:
                                    content = file_in_archive.read()
                                    print(f"Размер файла {file_name}: {len(content)} байт")
                                    
                                    if len(content) == 0:
                                        print(f"Файл {file_name} пустой")
                                        continue
                                    
                                    # Пробуем обработать как текстовый файл
                                    if file_name.lower().endswith(('.txt', '.md')):
                                        try:
                                            text = content.decode('utf-8')
                                            if text.strip():
                                                all_texts.append(f"=== ФАЙЛ В АРХИВЕ: {file_name} ===\n{text}")
                                                print(f"Успешно обработан текстовый файл: {file_name}")
                                                processed_files += 1
                                            else:
                                                print(f"Текстовый файл {file_name} пустой")
                                        except UnicodeDecodeError:
                                            try:
                                                text = content.decode('cp1251')
                                                if text.strip():
                                                    all_texts.append(f"=== ФАЙЛ В АРХИВЕ: {file_name} ===\n{text}")
                                                    print(f"Успешно обработан текстовый файл (cp1251): {file_name}")
                                                    processed_files += 1
                                                else:
                                                    print(f"Текстовый файл {file_name} пустой")
                                            except Exception as decode_error:
                                                print(f"Не удалось декодировать {file_name}: {str(decode_error)}")
                                    
                                    # Для других форматов сохраняем во временный файл и обрабатываем
                                    elif file_name.lower().endswith(('.pdf', '.docx', '.doc', '.rtf')):
                                        import tempfile
                                        
                                        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file_name).suffix) as temp_file:
                                            temp_file.write(content)
                                            temp_file_path = temp_file.name
                                        
                                        try:
                                            print(f"Временный файл создан: {temp_file_path}")
                                            
                                            # Специальная обработка для PDF файлов с OCR
                                            if file_name.lower().endswith('.pdf'):
                                                print(f"Обрабатываем PDF в архиве через OCR: {file_name}")
                                                try:
                                                    # Упрощенная обработка PDF через OCR для архивов
                                                    pdf_text = self._extract_from_pdf_simple(temp_file_path)
                                                    print(f"DEBUG: OCR результат для {file_name}: {len(pdf_text) if pdf_text else 0} символов")
                                                    if pdf_text and pdf_text.strip():
                                                        all_texts.append(f"=== ФАЙЛ В АРХИВЕ: {file_name} ===\n{pdf_text}")
                                                        print(f"✅ Успешно обработан PDF в архиве через OCR: {file_name}")
                                                        processed_files += 1
                                                    else:
                                                        print(f"❌ OCR не смог извлечь текст из PDF в архиве: {file_name}")
                                                        all_texts.append(f"=== ОШИБКА OCR PDF В АРХИВЕ: {file_name} ===\nOCR не смог распознать текст. Возможно, документ плохого качества или содержит только изображения.")
                                                        processed_files += 1  # Считаем как обработанный, даже если OCR не сработал
                                                except Exception as ocr_error:
                                                    print(f"❌ Ошибка OCR для PDF в архиве {file_name}: {str(ocr_error)}")
                                                    all_texts.append(f"=== ОШИБКА OCR PDF В АРХИВЕ: {file_name} ===\n{str(ocr_error)}")
                                                    processed_files += 1  # Считаем как обработанный, даже если произошла ошибка
                                            else:
                                                # Обычная обработка для других форматов
                                                text, success = self.extract_text(temp_file_path)
                                                if success and text.strip():
                                                    all_texts.append(f"=== ФАЙЛ В АРХИВЕ: {file_name} ===\n{text}")
                                                    print(f"Успешно обработан документ в архиве: {file_name}")
                                                    processed_files += 1
                                                else:
                                                    print(f"Не удалось извлечь текст из {file_name}")
                                                    all_texts.append(f"=== ОШИБКА ОБРАБОТКИ ФАЙЛА В АРХИВЕ: {file_name} ===\nНе удалось извлечь текст из документа.")
                                                    processed_files += 1  # Считаем как обработанный
                                                    
                                        except Exception as e:
                                            all_texts.append(f"=== ОШИБКА ОБРАБОТКИ ФАЙЛА В АРХИВЕ: {file_name} ===\n{str(e)}")
                                            print(f"Ошибка обработки {file_name}: {str(e)}")
                                            processed_files += 1  # Считаем как обработанный, даже если произошла ошибка
                                        finally:
                                            # Удаляем временный файл
                                            try:
                                                os.unlink(temp_file_path)
                                                print(f"Временный файл удален: {temp_file_path}")
                                            except Exception as cleanup_error:
                                                print(f"Ошибка удаления временного файла: {str(cleanup_error)}")
                                    
                            except Exception as e:
                                all_texts.append(f"=== ОШИБКА ЧТЕНИЯ ФАЙЛА В АРХИВЕ: {file_name} ===\n{str(e)}")
                                print(f"Ошибка чтения {file_name}: {str(e)}")
                                processed_files += 1  # Считаем как обработанный, даже если произошла ошибка
                    
                    if not found_supported_files:
                        raise Exception("В ZIP архиве нет поддерживаемых файлов")
                    
                    print(f"Обработано файлов в архиве: {processed_files}")
            
            elif ext == '.rar':
                # Для RAR файлов нужна дополнительная библиотека
                all_texts.append("RAR архивы пока не поддерживаются. Установите библиотеку rarfile для поддержки RAR.")
            
            result = "\n\n".join(all_texts) if all_texts else ""
            print(f"Итоговый результат из архива: {len(result)} символов")
            print(f"DEBUG: all_texts содержит {len(all_texts)} элементов")
            for i, text in enumerate(all_texts):
                print(f"DEBUG: элемент {i}: {len(text)} символов")
            
            # Если есть какие-то результаты (даже ошибки), возвращаем их
            if result.strip():
                print(f"DEBUG: Возвращаем результат длиной {len(result)} символов")
                return result
            else:
                print(f"DEBUG: Результат пустой, выбрасываем исключение")
                raise Exception("Не удалось извлечь текст из архива. Возможные причины: архив пустой, поврежден, содержит только неподдерживаемые файлы, или PDF файлы внутри архива требуют OCR.")
            
        except Exception as e:
            raise Exception(f"Ошибка обработки архива {filepath}: {str(e)}")

    def diagnose_file_problems(self, filepath: str) -> Dict[str, any]:
        """Детальная диагностика проблем с файлом"""
        diagnosis = {
            'filepath': filepath,
            'exists': False,
            'size': 0,
            'extension': '',
            'problems': [],
            'suggestions': []
        }
        
        try:
            # Проверяем существование файла
            if not os.path.exists(filepath):
                diagnosis['problems'].append("Файл не существует")
                diagnosis['suggestions'].append("Проверьте правильность пути к файлу")
                return diagnosis
            
            diagnosis['exists'] = True
            
            # Проверяем размер файла
            try:
                file_size = os.path.getsize(filepath)
                diagnosis['size'] = file_size
                if file_size == 0:
                    diagnosis['problems'].append("Файл пустой (0 байт)")
                    diagnosis['suggestions'].append("Файл может быть поврежден или не был полностью загружен")
            except Exception as e:
                diagnosis['problems'].append(f"Не удалось получить размер файла: {str(e)}")
            
            # Определяем расширение
            ext = Path(filepath).suffix.lower()
            diagnosis['extension'] = ext
            
            # Специфичные проверки по типам файлов
            if ext == '.pdf':
                diagnosis.update(self._diagnose_pdf(filepath))
            elif ext == '.zip':
                diagnosis.update(self._diagnose_zip(filepath))
            elif ext in ('.doc', '.rtf'):
                diagnosis.update(self._diagnose_legacy_doc(filepath))
            elif ext in self.SUPPORTED_EXTENSIONS['images']:
                diagnosis.update(self._diagnose_image(filepath))
            
        except Exception as e:
            diagnosis['problems'].append(f"Ошибка диагностики: {str(e)}")
        
        return diagnosis
    
    def _diagnose_pdf(self, filepath: str) -> Dict[str, any]:
        """Диагностика PDF файла"""
        diagnosis = {'problems': [], 'suggestions': []}
        
        try:
            # Проверяем через pdfplumber
            try:
                with pdfplumber.open(filepath) as pdf:
                    if len(pdf.pages) == 0:
                        diagnosis['problems'].append("PDF не содержит страниц")
                        diagnosis['suggestions'].append("Файл может быть поврежден")
                    
                    # Проверяем защиту
                    if hasattr(pdf, 'metadata') and pdf.metadata:
                        if 'Encrypt' in pdf.metadata:
                            diagnosis['problems'].append("PDF файл зашифрован")
                            diagnosis['suggestions'].append("Требуется пароль для открытия")
                    
                    # Проверяем наличие текста и изображений
                    has_text = False
                    has_images = False
                    total_images = 0
                    
                    for i, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        if text and text.strip():
                            has_text = True
                        if page.images:
                            has_images = True
                            total_images += len(page.images)
                    
                    if not has_text and has_images:
                        diagnosis['problems'].append("PDF содержит только изображения (отсканированный документ)")
                        diagnosis['suggestions'].append("Требуется OCR для извлечения текста")
                        if not TESSERACT_AVAILABLE:
                            diagnosis['problems'].append("Tesseract OCR не установлен")
                            diagnosis['suggestions'].append("Установите Tesseract OCR для распознавания отсканированных документов")
                            diagnosis['suggestions'].append("Запустите файл 'установить_tesseract.bat' для инструкций")
                        else:
                            # Проверяем poppler
                            try:
                                from pdf2image import convert_from_path
                                # Проверяем, установлен ли poppler
                                poppler_paths = [
                                    r"C:\poppler\bin",
                                    r"C:\Program Files\poppler\bin", 
                                    r"C:\Program Files (x86)\poppler\bin",
                                    os.path.join(os.path.expanduser("~"), "poppler", "bin"),
                                    r"C:\Users\Home\poppler-24.08.0\Library\bin"
                                ]
                                
                                poppler_found = False
                                for poppler_path in poppler_paths:
                                    if os.path.exists(os.path.join(poppler_path, "pdftoppm.exe")):
                                        poppler_found = True
                                        break
                                
                                if not poppler_found:
                                    # Проверяем через команду
                                    try:
                                        import subprocess
                                        result = subprocess.run(['pdftoppm', '-h'], 
                                                              capture_output=True, text=True, timeout=5)
                                        if result.returncode == 0:
                                            poppler_found = True
                                    except:
                                        pass
                                
                                if poppler_found:
                                    diagnosis['suggestions'].append("Tesseract и Poppler установлены - OCR должен работать")
                                else:
                                    diagnosis['problems'].append("Poppler не установлен или не найден в PATH")
                                    diagnosis['suggestions'].append("Установите poppler для конвертации PDF в изображения")
                                    diagnosis['suggestions'].append("Запустите файл 'установить_poppler_авто.bat' для автоматической установки")
                                    
                            except Exception as e:
                                diagnosis['suggestions'].append("Tesseract установлен, но есть проблемы с конвертацией PDF")
                    elif not has_text and not has_images:
                        diagnosis['problems'].append("PDF не содержит текста или изображений")
                        diagnosis['suggestions'].append("Файл может быть пустым или поврежденным")
                    elif has_text and has_images:
                        diagnosis['suggestions'].append("PDF содержит текст и изображения - должно обрабатываться нормально")
                        
            except Exception as e:
                error_msg = str(e).lower()
                if "password" in error_msg or "encrypted" in error_msg:
                    diagnosis['problems'].append("PDF файл защищен паролем")
                    diagnosis['suggestions'].append("Требуется пароль для открытия")
                elif "damaged" in error_msg or "corrupt" in error_msg:
                    diagnosis['problems'].append("PDF файл поврежден")
                    diagnosis['suggestions'].append("Попробуйте восстановить файл или получить новую копию")
                else:
                    diagnosis['problems'].append(f"Ошибка открытия PDF: {str(e)}")
                    diagnosis['suggestions'].append("Проверьте целостность файла")
                    
        except Exception as e:
            diagnosis['problems'].append(f"Ошибка диагностики PDF: {str(e)}")
        
        return diagnosis
    
    def _diagnose_zip(self, filepath: str) -> Dict[str, any]:
        """Диагностика ZIP файла"""
        diagnosis = {'problems': [], 'suggestions': []}
        
        try:
            with zipfile.ZipFile(filepath, 'r') as zip_file:
                # Проверяем целостность
                try:
                    zip_file.testzip()
                except Exception as e:
                    diagnosis['problems'].append("ZIP файл поврежден")
                    diagnosis['suggestions'].append("Попробуйте восстановить архив или получить новую копию")
                    return diagnosis
                
                # Проверяем содержимое
                file_list = zip_file.namelist()
                if not file_list:
                    diagnosis['problems'].append("ZIP архив пустой")
                    diagnosis['suggestions'].append("Архив не содержит файлов")
                    return diagnosis
                
                # Проверяем поддерживаемые файлы
                supported_extensions = ('.txt', '.md', '.pdf', '.docx', '.doc', '.rtf')
                supported_files = [f for f in file_list if any(f.lower().endswith(ext) for ext in supported_extensions)]
                
                if not supported_files:
                    diagnosis['problems'].append("ZIP не содержит поддерживаемых файлов")
                    diagnosis['suggestions'].append(f"Поддерживаемые форматы: {', '.join(supported_extensions)}")
                else:
                    diagnosis['suggestions'].append(f"Найдено {len(supported_files)} поддерживаемых файлов из {len(file_list)}")
                    
        except Exception as e:
            error_msg = str(e).lower()
            if "bad" in error_msg or "damaged" in error_msg:
                diagnosis['problems'].append("ZIP файл поврежден")
                diagnosis['suggestions'].append("Попробуйте восстановить архив")
            elif "password" in error_msg:
                diagnosis['problems'].append("ZIP файл защищен паролем")
                diagnosis['suggestions'].append("Требуется пароль для открытия")
            else:
                diagnosis['problems'].append(f"Ошибка открытия ZIP: {str(e)}")
        
        return diagnosis
    
    def _diagnose_legacy_doc(self, filepath: str) -> Dict[str, any]:
        """Диагностика старых форматов документов"""
        diagnosis = {'problems': [], 'suggestions': []}
        
        if not TEXTTRACT_AVAILABLE:
            diagnosis['problems'].append("Библиотека textract не установлена")
            diagnosis['suggestions'].append("Установите textract: pip install textract")
            diagnosis['suggestions'].append("Примечание: textract может не работать с Python 3.13")
        
        return diagnosis
    
    def _diagnose_image(self, filepath: str) -> Dict[str, any]:
        """Диагностика изображения"""
        diagnosis = {'problems': [], 'suggestions': []}
        
        if not TESSERACT_AVAILABLE:
            diagnosis['problems'].append("Tesseract OCR не установлен")
            diagnosis['suggestions'].append("Установите Tesseract OCR для распознавания текста")
        
        try:
            # Проверяем, можно ли открыть изображение
            with Image.open(filepath) as img:
                if img.size[0] < 100 or img.size[1] < 100:
                    diagnosis['problems'].append("Изображение слишком маленькое")
                    diagnosis['suggestions'].append("Для лучшего OCR используйте изображения с разрешением не менее 300 DPI")
        except Exception as e:
            diagnosis['problems'].append(f"Не удалось открыть изображение: {str(e)}")
            diagnosis['suggestions'].append("Проверьте формат и целостность файла")
        
        return diagnosis

class DocumentAnalyzerGUI:
    """Главный класс GUI приложения"""
    
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Анализатор документов v2.0")
        self.root.geometry("800x600")
        
        # Инициализация компонентов
        self.cache_manager = CacheManager()
        self.processor = DocumentProcessor(self.cache_manager)
        self.error_logger = ErrorLogger()  # Система логирования ошибок
        
        # Переменные состояния
        self.selected_folders = []
        
        # Умная настройка фильтров в зависимости от доступности компонентов
        self.file_filters = {
            'images': tk.BooleanVar(value=TESSERACT_AVAILABLE),  # Только если есть OCR
            'documents': tk.BooleanVar(value=True),  # DOCX всегда доступен
            'pdfs': tk.BooleanVar(value=True),
            'text': tk.BooleanVar(value=True),
            'archives': tk.BooleanVar(value=True)
        }
        self.processing = False
        self.cancel_processing = False
        
        # Результаты
        self.processed_files = []
        self.failed_files = []
        self.full_text_parts = []
        self.statistics = {}
        
        # Переменные для отслеживания времени
        self.start_time = None
        self.last_update_time = None
        self.estimated_total_time = None
        self.time_label = None
        
        self.setup_ui()
        self.check_dependencies()
    
    def setup_ui(self):
        """Настройка пользовательского интерфейса"""
        # Главный фрейм
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Настройка весов
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        
        # Заголовок
        title_label = ttk.Label(main_frame, text="Анализатор документов", 
                               font=('Arial', 16, 'bold'))
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 20))
        
        # Секция выбора папок
        folder_frame = ttk.LabelFrame(main_frame, text="Выбор папки", padding="10")
        folder_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        folder_frame.columnconfigure(1, weight=1)
        
        ttk.Button(folder_frame, text="Выбрать папку", 
                  command=self.select_folders).grid(row=0, column=0, padx=(0, 10))
        
        self.folder_label = ttk.Label(folder_frame, text="Папки не выбраны")
        self.folder_label.grid(row=0, column=1, sticky=(tk.W, tk.E))
        
        # Секция фильтров
        filter_frame = ttk.LabelFrame(main_frame, text="Фильтры файлов", padding="10")
        filter_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        filters = [
            ('Изображения (JPG, PNG, BMP)', 'images', TESSERACT_AVAILABLE),
            ('Документы (DOCX, DOC, RTF)', 'documents', True),
            ('PDF файлы', 'pdfs', True),
            ('Текстовые файлы (TXT, MD)', 'text', True),
            ('Архивы (ZIP, RAR)', 'archives', True)
        ]
        
        for i, (label, key, available) in enumerate(filters):
            if not available:
                label += " (недоступно)"
            checkbox = ttk.Checkbutton(filter_frame, text=label, 
                                      variable=self.file_filters[key],
                                      state='normal' if available else 'disabled')
            checkbox.grid(row=i//2, column=i%2, sticky=tk.W, padx=5)
        
        # Секция управления
        control_frame = ttk.Frame(main_frame)
        control_frame.grid(row=3, column=0, columnspan=2, pady=10)
        
        self.start_button = ttk.Button(control_frame, text="Начать анализ", 
                                      command=self.start_processing)
        self.start_button.pack(side=tk.LEFT, padx=(0, 10))
        
        self.cancel_button = ttk.Button(control_frame, text="Отменить", 
                                       command=self.cancel_processing_func, state=tk.DISABLED)
        self.cancel_button.pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(control_frame, text="Очистить кэш", 
                  command=self.clear_cache).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(control_frame, text="Диагностика ошибок", 
                  command=self.diagnose_failed_files).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(control_frame, text="Просмотр логов", 
                  command=self.show_logs).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(control_frame, text="Сохранить результаты", 
                  command=self.save_results).pack(side=tk.LEFT)
        
        # Прогресс
        progress_frame = ttk.Frame(main_frame)
        progress_frame.grid(row=4, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        progress_frame.columnconfigure(0, weight=1)
        
        self.progress_var = tk.StringVar(value="Готов к работе")
        ttk.Label(progress_frame, textvariable=self.progress_var).grid(row=0, column=0, sticky=tk.W)
        
        self.progress_bar = ttk.Progressbar(progress_frame, mode='determinate')
        self.progress_bar.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(5, 0))
        
        # Время обработки
        self.time_var = tk.StringVar(value="")
        self.time_label = ttk.Label(progress_frame, textvariable=self.time_var, font=('Arial', 9))
        self.time_label.grid(row=2, column=0, sticky=tk.W, pady=(2, 0))
        
        # Лог
        log_frame = ttk.LabelFrame(main_frame, text="Лог операций", padding="10")
        log_frame.grid(row=5, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)
        main_frame.rowconfigure(5, weight=1)
        
        self.log_text = ScrolledText(log_frame, height=10, wrap=tk.WORD)
        self.log_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Статистика
        stats_frame = ttk.LabelFrame(main_frame, text="Статистика", padding="10")
        stats_frame.grid(row=6, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        self.stats_label = ttk.Label(stats_frame, text="Нет данных")
        self.stats_label.pack()
    
    def check_dependencies(self):
        """Проверка зависимостей"""
        missing = DependencyChecker.get_missing_dependencies()
        
        # Проверяем только критически важные зависимости
        critical_missing = []
        optional_missing = []
        
        for dep in missing:
            if "Tesseract OCR" in dep or "pytesseract" in dep:
                critical_missing.append(dep)
            elif "textract" in dep:
                optional_missing.append(dep)
            else:
                critical_missing.append(dep)
        
        if critical_missing:
            self.log_message("⚠️ Отсутствуют важные компоненты:")
            for dep in critical_missing:
                self.log_message(f"  - {dep}")
            self.log_message("📋 Для полной функциональности установите Tesseract OCR")
        elif optional_missing:
            self.log_message("ℹ️ Дополнительные возможности:")
            for dep in optional_missing:
                self.log_message(f"  - {dep} (для старых форматов .doc/.rtf)")
            self.log_message("📋 Эти компоненты необязательны, программа работает без них")
        else:
            self.log_message("✅ Все компоненты установлены")
    
    def log_message(self, message: str, level: str = "INFO"):
        """Добавление сообщения в лог"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()
        
        # Записываем в файл лога сессии
        self.error_logger.log_session(message, level)
    
    def select_folders(self):
        """Выбор папок для анализа"""
        try:
            self.log_message("🔄 Открытие диалога выбора папок...")
            
            # Обновляем интерфейс
            self.root.update()
            self.root.update_idletasks()
            
            # Принудительно выводим окно на передний план
            self.root.lift()
            self.root.attributes('-topmost', True)
            self.root.focus_force()
            
            # Небольшая задержка для стабильности
            self.root.after(100)
            
            # Пробуем разные варианты диалога
            folders = None
            
            # Вариант 1: Обычный диалог
            try:
                folders = filedialog.askdirectory(
                    title="Выберите папку для анализа", 
                    initialdir=os.path.expanduser("~")
                )
                if folders:
                    folders = (folders,)
            except Exception as e1:
                self.log_message(f"Вариант 1 не сработал: {e1}")
                
                # Вариант 2: Диалог без multiple
                try:
                    folder = filedialog.askdirectory(
                        title="Выберите папку для анализа",
                        initialdir=os.path.expanduser("~")
                    )
                    if folder:
                        folders = (folder,)
                except Exception as e2:
                    self.log_message(f"Вариант 2 не сработал: {e2}")
                    
                    # Вариант 3: Простой диалог
                    try:
                        folders = filedialog.askdirectory(
                            title="Выберите папку",
                            initialdir="C:\\"
                        )
                        if folders:
                            folders = (folders,)
                    except Exception as e3:
                        self.log_message(f"Вариант 3 не сработал: {e3}")
                        raise Exception(f"Все варианты диалога не работают: {e1}, {e2}, {e3}")
            
            # Возвращаем окно в нормальное состояние
            self.root.attributes('-topmost', False)
            
            if folders:
                self.selected_folders = list(folders)
                folder_names = [os.path.basename(f) for f in self.selected_folders]
                self.folder_label.config(text=f"Выбрано: {', '.join(folder_names)}")
                self.log_message(f"✅ Выбрано папок: {len(self.selected_folders)}")
                self.log_message(f"📁 Папки: {', '.join(folder_names)}")
            else:
                self.log_message("ℹ️ Папки не выбраны")
                
        except Exception as e:
            self.log_message(f"❌ Ошибка при выборе папок: {e}")
            messagebox.showerror("Ошибка", f"Не удалось открыть диалог выбора папок:\n{str(e)}")
        finally:
            # Убираем topmost в любом случае
            try:
                self.root.attributes('-topmost', False)
            except:
                pass
    

    
    def get_filtered_files(self) -> List[str]:
        """Получение отфильтрованного списка файлов"""
        all_files = []
        
        for folder in self.selected_folders:
            for root, dirs, files in os.walk(folder):
                for file in files:
                    filepath = os.path.join(root, file)
                    ext = Path(filepath).suffix.lower()
                    
                    # Проверяем фильтры
                    include_file = False
                    for filter_type, var in self.file_filters.items():
                        if var.get() and ext in self.processor.SUPPORTED_EXTENSIONS[filter_type]:
                            include_file = True
                            break
                    
                    if include_file:
                        all_files.append(filepath)
        
        return all_files
    
    def start_processing(self):
        """Запуск обработки файлов"""
        if not self.selected_folders:
            messagebox.showwarning("Предупреждение", "Выберите папки для анализа")
            return
        
        if self.processing:
            return
        
        self.processing = True
        self.cancel_processing = False
        
        # Сбрасываем флаг отмены в процессоре
        if hasattr(self, 'processor'):
            self.processor.set_cancelled(False)
        
        self.start_button.config(state=tk.DISABLED)
        self.cancel_button.config(state=tk.NORMAL)
        
        # Очищаем предыдущие результаты
        self.processed_files = []
        self.failed_files = []
        self.full_text_parts = []
        self.statistics = {}
        
        # Переменные для отслеживания времени
        self.start_time = time.time()
        self.last_update_time = self.start_time
        self.estimated_total_time = None
        self.time_label = None
        
        # Запускаем обработку в отдельном потоке
        thread = threading.Thread(target=self.process_files_thread)
        thread.daemon = True
        thread.start()
    
    def process_files_thread(self):
        """Обработка файлов в отдельном потоке"""
        try:
            # Получаем список файлов
            all_files = self.get_filtered_files()
            if not all_files:
                self.log_message("❌ Не найдено файлов для обработки")
                return
            
            # Инициализируем отслеживание времени
            self.start_time = time.time()
            self.last_update_time = self.start_time
            
            total_files = len(all_files)
            self.log_message(f"🚀 Начинаем обработку {total_files} файлов...")
            
            # Настраиваем прогресс-бар
            self.progress_bar['maximum'] = total_files
            self.progress_bar['value'] = 0
            
            processed_count = 0
            failed_count = 0
            
            for i, filepath in enumerate(all_files):
                if self.cancel_processing:
                    self.log_message("⏹️ Обработка отменена пользователем")
                    break
                
                # Обновляем прогресс
                self.progress_var.set(f"Обрабатываем: {i+1}/{total_files} - {Path(filepath).name}")
                self.progress_bar['value'] = i + 1
                
                # Обновляем оценку времени каждые 5 файлов или для каждого файла после 10-го
                if i % 5 == 0 or i >= 10:
                    self.update_time_estimate(i + 1, total_files)
                
                # Обрабатываем файл
                try:
                    result = self.process_single_file(filepath)
                    if result:
                        self.processed_files.append(filepath)
                        processed_count += 1
                        self.log_message(f"✅ {Path(filepath).name}")
                    else:
                        self.failed_files.append(filepath)
                        failed_count += 1
                        self.log_message(f"❌ {Path(filepath).name}")
                except Exception as e:
                    self.failed_files.append(filepath)
                    failed_count += 1
                    self.log_message(f"❌ {Path(filepath).name} - {str(e)}")
                    self.error_logger.log_error(filepath, str(e))
            
            # Завершаем обработку
            self.finalize_processing()
            
        except Exception as e:
            self.log_message(f"❌ Критическая ошибка: {str(e)}")
            self.error_logger.log_error("", f"Критическая ошибка: {str(e)}")
        finally:
            # Сбрасываем отслеживание времени
            self.reset_time_tracking()
            self.processing = False
            self.start_button.config(state=tk.NORMAL)
            self.cancel_button.config(state=tk.DISABLED)
    
    def process_single_file(self, filepath: str) -> Optional[str]:
        """Обработка одного файла"""
        try:
            text, success = self.processor.extract_text(filepath)
            
            if success and text.strip():
                relative_path = os.path.relpath(filepath, self.selected_folders[0])
                self.log_message(f"✅ {relative_path}")
                
                # Добавляем в общий текст
                part = f"{'='*60}\nФАЙЛ: {relative_path}\n{'='*60}\n{text.strip()}\n"
                self.full_text_parts.append(part)
                
                # Обновляем статистику
                self.update_statistics(filepath, len(text))
                
                return relative_path
            else:
                relative_path = os.path.relpath(filepath, self.selected_folders[0])
                self.log_message(f"❌ {relative_path}")
                self.failed_files.append(relative_path)
                
                # Записываем ошибку в лог
                self.error_logger.log_error(
                    filepath, 
                    "Не удалось извлечь текст или файл пустой", 
                    "ОБРАБОТКА"
                )
                
        except Exception as e:
            relative_path = os.path.relpath(filepath, self.selected_folders[0])
            self.log_message(f"❌ {relative_path}: {e}", "ERROR")
            self.failed_files.append(relative_path)
            
            # Записываем ошибку в лог
            self.error_logger.log_error(
                filepath, 
                str(e), 
                "ИСКЛЮЧЕНИЕ"
            )
        
        return None
    
    def update_statistics(self, filepath: str, text_length: int):
        """Обновление статистики"""
        ext = Path(filepath).suffix.lower()
        
        if ext not in self.statistics:
            self.statistics[ext] = {'count': 0, 'total_chars': 0}
        
        self.statistics[ext]['count'] += 1
        self.statistics[ext]['total_chars'] += text_length
    
    def finalize_processing(self):
        """Завершение обработки"""
        if self.cancel_processing:
            self.log_message("⏹️ Обработка отменена")
            self.progress_var.set("Обработка отменена")
            return
        
        # Рассчитываем итоговое время
        if self.start_time:
            total_time = time.time() - self.start_time
            if total_time > 3600:
                hours = int(total_time // 3600)
                minutes = int((total_time % 3600) // 60)
                time_str = f"{hours}ч {minutes}мин"
            elif total_time > 60:
                minutes = int(total_time // 60)
                seconds = int(total_time % 60)
                time_str = f"{minutes}мин {seconds}сек"
            else:
                seconds = int(total_time)
                time_str = f"{seconds}сек"
        else:
            time_str = "неизвестно"
        
        # Показываем результаты
        total_files = len(self.processed_files) + len(self.failed_files)
        success_rate = (len(self.processed_files) / total_files * 100) if total_files > 0 else 0
        
        self.log_message(f"🎉 Обработка завершена за {time_str}!")
        self.log_message(f"📊 Результаты: {len(self.processed_files)} успешно, {len(self.failed_files)} ошибок ({success_rate:.1f}%)")
        
        # Обновляем прогресс
        self.progress_var.set(f"Завершено: {len(self.processed_files)}/{total_files} файлов за {time_str}")
        self.progress_bar['value'] = total_files
        
        # Сбрасываем время
        self.reset_time_tracking()
        
        # Показываем сообщение о завершении
        if len(self.failed_files) > 0:
            result = messagebox.askyesno("Обработка завершена", 
                              f"Обработано {len(self.processed_files)} из {total_files} файлов.\n"
                              f"Время: {time_str}\n"
                              f"Ошибок: {len(self.failed_files)}\n\n"
                              f"Сохранить результаты анализа?")
        else:
            result = messagebox.askyesno("Обработка завершена", 
                              f"Все {total_files} файлов обработаны успешно!\n"
                              f"Время: {time_str}\n\n"
                              f"Сохранить результаты анализа?")
        
        # Автоматически предлагаем сохранить результаты
        if result and self.full_text_parts:
            self.save_results()
    
    def save_results(self):
        """Сохранение результатов"""
        if not self.full_text_parts:
            messagebox.showwarning("Предупреждение", "Нет результатов для сохранения. Сначала запустите анализ документов.")
            return
        
        # Прямо открываем диалог сохранения с выбором формата
        filetypes = [
            ('Text Files', '*.txt'),
            ('Markdown', '*.md'),
            ('Word Document', '*.docx')
        ]
        
        path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=filetypes,
            title="Сохранить результат анализа"
        )
        
        if path:
            try:
                # Определяем формат по расширению файла
                ext = Path(path).suffix.lower()
                
                if ext == '.docx':
                    doc = Document()
                    for part in self.full_text_parts:
                        doc.add_paragraph(part)
                    doc.save(path)
                else:
                    with open(path, 'w', encoding='utf-8') as f:
                        f.write("\n".join(self.full_text_parts))
                
                self.log_message(f"✅ Результат сохранён: {path}")
                
                # Сохраняем кэш
                self.cache_manager.save_cache()
                
                # Показываем сообщение об успехе
                messagebox.showinfo("Сохранение", f"Результат успешно сохранён в:\n{path}")
                
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить файл:\n{str(e)}")
    
    def cancel_processing_func(self):
        """Отмена обработки"""
        self.cancel_processing = True
        if hasattr(self, 'processor'):
            self.processor.set_cancelled(True)
        self.log_message("Отмена обработки...")
    
    def clear_cache(self):
        """Очистка кэша"""
        self.cache_manager.clear_cache()
        self.log_message("Кэш очищен")
    
    def diagnose_failed_files(self):
        """Диагностика файлов, которые не удалось обработать"""
        if not hasattr(self, 'failed_files') or not self.failed_files:
            messagebox.showinfo("Диагностика", "Нет файлов для диагностики. Сначала запустите анализ.")
            return
        
        # Создаем окно диагностики
        diagnosis_window = tk.Toplevel(self.root)
        diagnosis_window.title("Диагностика проблем с файлами")
        diagnosis_window.geometry("900x700")
        
        # Создаем notebook для вкладок
        notebook = ttk.Notebook(diagnosis_window)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Вкладка с общей сводкой
        summary_frame = ttk.Frame(notebook)
        notebook.add(summary_frame, text="Сводка")
        
        summary_text = ScrolledText(summary_frame, wrap=tk.WORD)
        summary_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Вкладка с детальной диагностикой
        details_frame = ttk.Frame(notebook)
        notebook.add(details_frame, text="Детали")
        
        details_text = ScrolledText(details_frame, wrap=tk.WORD)
        details_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Кнопки управления
        button_frame = ttk.Frame(diagnosis_window)
        button_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        def run_diagnosis():
            summary_text.delete(1.0, tk.END)
            details_text.delete(1.0, tk.END)
            
            summary_text.insert(tk.END, "🔍 Запуск диагностики...\n\n")
            details_text.insert(tk.END, "🔍 Детальная диагностика...\n\n")
            
            # Получаем полные пути к файлам
            full_paths = []
            for failed_file in self.failed_files:
                for folder in self.selected_folders:
                    full_path = os.path.join(folder, failed_file)
                    if os.path.exists(full_path):
                        full_paths.append(full_path)
                        break
            
            if not full_paths:
                summary_text.insert(tk.END, "❌ Не удалось найти файлы для диагностики\n")
                return
            
            summary_text.insert(tk.END, f"📁 Найдено {len(full_paths)} файлов для диагностики\n\n")
            
            # Группируем проблемы
            problem_groups = {}
            total_problems = 0
            
            for i, filepath in enumerate(full_paths):
                summary_text.insert(tk.END, f"🔍 Диагностика {i+1}/{len(full_paths)}: {os.path.basename(filepath)}\n")
                summary_text.see(tk.END)
                diagnosis_window.update()
                
                # Запускаем диагностику
                diagnosis = self.processor.diagnose_file_problems(filepath)
                
                # Добавляем в детали
                details_text.insert(tk.END, f"{'='*60}\n")
                details_text.insert(tk.END, f"ФАЙЛ: {os.path.basename(filepath)}\n")
                details_text.insert(tk.END, f"ПУТЬ: {filepath}\n")
                details_text.insert(tk.END, f"РАЗМЕР: {diagnosis['size']} байт\n")
                details_text.insert(tk.END, f"РАСШИРЕНИЕ: {diagnosis['extension']}\n\n")
                
                if diagnosis['problems']:
                    details_text.insert(tk.END, "❌ ПРОБЛЕМЫ:\n")
                    for problem in diagnosis['problems']:
                        details_text.insert(tk.END, f"  • {problem}\n")
                        # Группируем проблемы
                        problem_type = problem.split(':')[0] if ':' in problem else problem
                        problem_groups[problem_type] = problem_groups.get(problem_type, 0) + 1
                        total_problems += 1
                    
                    details_text.insert(tk.END, "\n💡 РЕКОМЕНДАЦИИ:\n")
                    for suggestion in diagnosis['suggestions']:
                        details_text.insert(tk.END, f"  • {suggestion}\n")
                else:
                    details_text.insert(tk.END, "✅ Проблем не обнаружено\n")
                
                details_text.insert(tk.END, "\n")
                details_text.see(tk.END)
            
            # Обновляем сводку
            summary_text.insert(tk.END, f"\n📊 СВОДКА ДИАГНОСТИКИ:\n")
            summary_text.insert(tk.END, f"Всего файлов: {len(full_paths)}\n")
            summary_text.insert(tk.END, f"Всего проблем: {total_problems}\n\n")
            
            if problem_groups:
                summary_text.insert(tk.END, "📈 ТИПЫ ПРОБЛЕМ:\n")
                for problem_type, count in sorted(problem_groups.items(), key=lambda x: x[1], reverse=True):
                    summary_text.insert(tk.END, f"  • {problem_type}: {count} файлов\n")
                
                summary_text.insert(tk.END, f"\n💡 ОБЩИЕ РЕКОМЕНДАЦИИ:\n")
                if "PDF файл зашифрован" in problem_groups:
                    summary_text.insert(tk.END, "  • Многие PDF файлы защищены паролем. Требуется пароль для их обработки.\n")
                if "PDF содержит только изображения" in problem_groups:
                    summary_text.insert(tk.END, "  • Некоторые PDF содержат только изображения. Убедитесь, что установлен Tesseract OCR.\n")
                if "ZIP файл поврежден" in problem_groups:
                    summary_text.insert(tk.END, "  • Некоторые ZIP архивы повреждены. Попробуйте восстановить их.\n")
                if "ZIP не содержит поддерживаемых файлов" in problem_groups:
                    summary_text.insert(tk.END, "  • Некоторые архивы не содержат поддерживаемых форматов файлов.\n")
            else:
                summary_text.insert(tk.END, "✅ Все файлы в порядке!\n")
        
        def export_diagnosis():
            try:
                filepath = filedialog.asksaveasfilename(
                    defaultextension=".txt",
                    filetypes=[('Text Files', '*.txt')],
                    title="Сохранить диагностику"
                )
                if filepath:
                    with open(filepath, 'w', encoding='utf-8') as f:
                        f.write("ДИАГНОСТИКА ПРОБЛЕМ С ФАЙЛАМИ\n")
                        f.write("="*50 + "\n\n")
                        f.write(summary_text.get(1.0, tk.END))
                        f.write("\n" + "="*50 + "\n\n")
                        f.write("ДЕТАЛЬНАЯ ДИАГНОСТИКА\n")
                        f.write("="*50 + "\n\n")
                        f.write(details_text.get(1.0, tk.END))
                    
                    messagebox.showinfo("Экспорт", f"Диагностика сохранена в: {filepath}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить диагностику: {e}")
        
        ttk.Button(button_frame, text="Запустить диагностику", 
                  command=run_diagnosis).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(button_frame, text="Экспорт диагностики", 
                  command=export_diagnosis).pack(side=tk.LEFT)
        
        # Запускаем диагностику автоматически
        diagnosis_window.after(100, run_diagnosis)
    
    def show_logs(self):
        """Показать окно с логами"""
        log_window = tk.Toplevel(self.root)
        log_window.title("Логи ошибок и сессии")
        log_window.geometry("800x600")
        
        # Создаём notebook для вкладок
        notebook = ttk.Notebook(log_window)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Вкладка с ошибками
        error_frame = ttk.Frame(notebook)
        notebook.add(error_frame, text="Ошибки")
        
        error_text = ScrolledText(error_frame, wrap=tk.WORD)
        error_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Загружаем содержимое лога ошибок
        try:
            if self.error_logger.error_log_file.exists():
                with open(self.error_logger.error_log_file, 'r', encoding='utf-8') as f:
                    error_text.insert(tk.END, f.read())
            else:
                error_text.insert(tk.END, "Лог ошибок пуст")
        except Exception as e:
            error_text.insert(tk.END, f"Ошибка чтения лога: {e}")
        
        # Вкладка с сессией
        session_frame = ttk.Frame(notebook)
        notebook.add(session_frame, text="Сессия")
        
        session_text = ScrolledText(session_frame, wrap=tk.WORD)
        session_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Загружаем содержимое лога сессии
        try:
            if self.error_logger.session_log_file.exists():
                with open(self.error_logger.session_log_file, 'r', encoding='utf-8') as f:
                    session_text.insert(tk.END, f.read())
            else:
                session_text.insert(tk.END, "Лог сессии пуст")
        except Exception as e:
            session_text.insert(tk.END, f"Ошибка чтения лога: {e}")
        
        # Кнопки управления
        button_frame = ttk.Frame(log_window)
        button_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        def export_errors():
            try:
                output_file = self.error_logger.export_errors_to_txt()
                if output_file:
                    messagebox.showinfo("Экспорт", f"Ошибки экспортированы в: {output_file}")
                else:
                    messagebox.showerror("Ошибка", "Не удалось экспортировать ошибки")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка экспорта: {e}")
        
        def open_log_folder():
            try:
                os.startfile(self.error_logger.log_dir)
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось открыть папку: {e}")
        
        ttk.Button(button_frame, text="Экспорт ошибок", 
                  command=export_errors).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(button_frame, text="Открыть папку логов", 
                  command=open_log_folder).pack(side=tk.LEFT)
    
    def run(self):
        """Запуск приложения"""
        self.root.mainloop()
    
    def update_time_estimate(self, current_file_index: int, total_files: int):
        """Обновляет оценку оставшегося времени"""
        if not self.start_time or current_file_index == 0:
            return
        
        current_time = time.time()
        elapsed_time = current_time - self.start_time
        
        # Рассчитываем среднее время на файл
        avg_time_per_file = elapsed_time / current_file_index
        
        # Рассчитываем оставшееся время
        remaining_files = total_files - current_file_index
        estimated_remaining_time = avg_time_per_file * remaining_files
        
        # Форматируем время
        if estimated_remaining_time > 3600:  # больше часа
            hours = int(estimated_remaining_time // 3600)
            minutes = int((estimated_remaining_time % 3600) // 60)
            time_str = f"⏱️ Осталось примерно: {hours}ч {minutes}мин"
        elif estimated_remaining_time > 60:  # больше минуты
            minutes = int(estimated_remaining_time // 60)
            seconds = int(estimated_remaining_time % 60)
            time_str = f"⏱️ Осталось примерно: {minutes}мин {seconds}сек"
        else:
            seconds = int(estimated_remaining_time)
            time_str = f"⏱️ Осталось примерно: {seconds}сек"
        
        # Добавляем информацию о прогрессе
        progress_percent = (current_file_index / total_files) * 100
        time_str += f" | Прогресс: {current_file_index}/{total_files} ({progress_percent:.1f}%)"
        
        # Обновляем метку времени
        if self.time_var:
            self.time_var.set(time_str)
        
        # Обновляем время последнего обновления
        self.last_update_time = current_time
    
    def reset_time_tracking(self):
        """Сбрасывает отслеживание времени"""
        self.start_time = None
        self.last_update_time = None
        self.estimated_total_time = None
        if self.time_var:
            self.time_var.set("")

def main():
    """Главная функция"""
    app = DocumentAnalyzerGUI()
    app.run()

if __name__ == "__main__":
    main() 